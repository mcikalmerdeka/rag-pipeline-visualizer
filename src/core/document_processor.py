"""
Document processing utilities for loading documents from various file formats.
Supports PDF, Word (DOCX/DOC), Images, and TXT files.
"""

import os
import io
import base64
from typing import List, Dict, Any, Optional
from datetime import datetime

from src.config import logger_text
from src.config.settings import SUPPORTED_FILE_TYPES


def get_file_extension(filename: str) -> str:
    """Extract file extension from filename"""
    return os.path.splitext(filename)[1].lower().lstrip('.')


def get_file_type(file_path: str) -> str:
    """Determine file type based on extension"""
    ext = get_file_extension(file_path)

    for file_type, extensions in SUPPORTED_FILE_TYPES.items():
        if ext in extensions:
            return file_type
    return "unknown"


def extract_metadata_from_uploaded_file(uploaded_file, file_type: str) -> Dict[str, Any]:
    """
    Extract metadata from an uploaded file object

    Args:
        uploaded_file: Streamlit uploaded file object
        file_type: Type of file (pdf, word, image, text)

    Returns:
        Dictionary containing file metadata
    """
    return {
        "source": uploaded_file.name,
        "filename": uploaded_file.name,
        "file_size_bytes": len(uploaded_file.getvalue()),
        "uploaded_at": datetime.now().isoformat(),
        "file_type": file_type.upper()
    }


def load_txt_from_memory(uploaded_file) -> str:
    """
    Load text file from uploaded file buffer (in-memory)

    Args:
        uploaded_file: Streamlit uploaded file object

    Returns:
        Text content as string
    """
    logger_text.info(f"Loading text document from memory: {uploaded_file.name}")
    content = uploaded_file.read().decode('utf-8')
    logger_text.info(f"Loaded {len(content)} characters from text file")
    return content


def load_pdf_from_memory(uploaded_file) -> str:
    """
    Load PDF document from uploaded file buffer (in-memory)

    Args:
        uploaded_file: Streamlit uploaded file object with getbuffer() method

    Returns:
        Extracted text content as string

    Raises:
        ValueError: If no content could be loaded from the PDF
    """
    logger_text.info(f"Loading PDF document from memory: {uploaded_file.name}")

    try:
        import tempfile

        # Write to temporary file for PyMuPDFLoader (it requires a file path)
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
            tmp_file.write(uploaded_file.getbuffer())
            tmp_path = tmp_file.name

        try:
            from langchain_community.document_loaders import PyMuPDFLoader

            document_loader = PyMuPDFLoader(tmp_path)
            docs = document_loader.load()

            if not docs:
                logger_text.error(f"Failed to load any content from PDF: {uploaded_file.name}")
                raise ValueError(f"Failed to load any content from PDF: {uploaded_file.name}")

            # Combine all pages into single text
            content = "\n\n".join([doc.page_content for doc in docs])
            logger_text.info(f"Loaded {len(docs)} pages ({len(content)} chars) from PDF")
            return content
        finally:
            # Clean up temporary file
            try:
                os.unlink(tmp_path)
            except Exception as e:
                logger_text.warning(f"Failed to delete temporary file {tmp_path}: {e}")

    except Exception as e:
        logger_text.error(f"Failed to load PDF from memory: {uploaded_file.name}, Error: {str(e)}")
        raise ValueError(f"Failed to load PDF: {str(e)}")


def load_word_from_memory(uploaded_file) -> str:
    """
    Load Word document from uploaded file buffer (in-memory)

    Args:
        uploaded_file: Streamlit uploaded file object

    Returns:
        Extracted text content as string

    Raises:
        ValueError: If no content could be loaded from the document
        ImportError: If required dependencies are not installed
    """
    logger_text.info(f"Loading Word document from memory: {uploaded_file.name}")

    try:
        from docx import Document as DocxDocument
    except ImportError:
        logger_text.error("python-docx not installed. Install with: pip install python-docx")
        raise ImportError(
            "python-docx is required for Word document processing. "
            "Install it with: pip install python-docx"
        )

    try:
        # Load directly from BytesIO buffer
        doc = DocxDocument(io.BytesIO(uploaded_file.getbuffer()))

        # Extract text from all paragraphs
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))

        content = "\n".join(full_text)

        if not content.strip():
            logger_text.error(f"No text content found in Word document: {uploaded_file.name}")
            raise ValueError(f"No text content found in Word document: {uploaded_file.name}")

        logger_text.info(f"Successfully loaded Word document ({len(content)} chars)")
        return content

    except Exception as e:
        logger_text.error(f"Failed to load Word document from memory: {uploaded_file.name}, Error: {str(e)}")
        raise ValueError(f"Failed to load Word document: {str(e)}")


def load_image_from_memory(uploaded_file) -> Dict[str, Any]:
    """
    Load image document from uploaded file buffer and extract metadata.
    For the RAG visualizer, images are converted to a descriptive text format.

    Args:
        uploaded_file: Streamlit uploaded file object

    Returns:
        Dictionary containing extracted text representation and metadata

    Raises:
        ValueError: If image cannot be processed
        ImportError: If required dependencies are not installed
    """
    logger_text.info(f"Processing image from memory: {uploaded_file.name}")

    try:
        from PIL import Image
    except ImportError as e:
        logger_text.error("Pillow not installed.")
        raise ImportError(
            "Pillow is required for image processing. "
            "Install with: pip install Pillow"
        )

    try:
        # Open image from BytesIO buffer to get dimensions
        image = Image.open(io.BytesIO(uploaded_file.getbuffer()))

        # Convert to RGB if necessary
        if image.mode != 'RGB':
            image = image.convert('RGB')

        # Encode image to base64 for potential multimodal use
        file_bytes = uploaded_file.getvalue()
        base64_encoded = base64.b64encode(file_bytes).decode('utf-8')

        # Determine MIME type from file extension
        ext = get_file_extension(uploaded_file.name)
        mime_types = {
            'png': 'image/png',
            'jpg': 'image/jpeg',
            'jpeg': 'image/jpeg',
            'gif': 'image/gif',
            'bmp': 'image/bmp',
            'tiff': 'image/tiff',
            'webp': 'image/webp'
        }
        mime_type = mime_types.get(ext, 'image/jpeg')
        base64_image = f"data:{mime_type};base64,{base64_encoded}"

        # Create descriptive text content for the image
        content = f"[IMAGE: {uploaded_file.name} - {image.width}x{image.height} pixels]"

        result = {
            "content": content,
            "metadata": {
                "source": uploaded_file.name,
                "filename": uploaded_file.name,
                "file_size_bytes": len(file_bytes),
                "uploaded_at": datetime.now().isoformat(),
                "file_type": "IMAGE",
                "image_width": image.width,
                "image_height": image.height,
                "image_mode": image.mode,
                "base64_image": base64_image,
                "is_multimodal": True
            }
        }

        logger_text.info(f"Successfully processed image: {uploaded_file.name} ({image.width}x{image.height})")
        return result

    except Exception as e:
        logger_text.error(f"Failed to process image from memory: {uploaded_file.name}, Error: {str(e)}")
        raise ValueError(f"Failed to process image: {str(e)}")


def load_document_from_memory(uploaded_file) -> str:
    """
    Load document from various file formats (PDF, Word, Image, TXT) directly from memory.
    Returns the text content as a string.

    Args:
        uploaded_file: Streamlit uploaded file object

    Returns:
        Text content extracted from the file

    Raises:
        ValueError: If file type is not supported or content cannot be loaded
    """
    file_type = get_file_type(uploaded_file.name)

    logger_text.info(f"Loading document of type '{file_type}' from memory: {uploaded_file.name}")

    if file_type == "text":
        return load_txt_from_memory(uploaded_file)
    elif file_type == "pdf":
        return load_pdf_from_memory(uploaded_file)
    elif file_type == "word":
        return load_word_from_memory(uploaded_file)
    elif file_type == "image":
        # For images, return the content string
        result = load_image_from_memory(uploaded_file)
        return result["content"]
    else:
        ext = get_file_extension(uploaded_file.name)
        logger_text.error(f"Unsupported file type: {ext}")
        raise ValueError(
            f"Unsupported file type: '.{ext}'. "
            f"Supported formats: PDF, Word (DOCX/DOC), Images (PNG, JPG, JPEG, GIF, BMP, TIFF, WEBP), TXT"
        )


def get_supported_file_types() -> List[str]:
    """Get list of supported file extensions for file uploader"""
    extensions = []
    for ext_list in SUPPORTED_FILE_TYPES.values():
        extensions.extend(ext_list)
    return extensions


def get_file_type_label(filename: str) -> str:
    """Get a human-readable label for the file type"""
    file_type = get_file_type(filename)
    labels = {
        'text': 'Text Document',
        'pdf': 'PDF Document',
        'word': 'Word Document',
        'image': 'Image File'
    }
    return labels.get(file_type, 'Unknown File Type')
